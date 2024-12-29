import os
from docx import Document
from django.conf import settings

def fill_clearance_template(template_path, output_path, data):
    """
    Fill the Barangay Clearance Word template with provided data.
    :param template_path: Path to the Word template.
    :param output_path: Path to save the filled document.
    :param data: Dictionary containing placeholders and their values.
    """
    try:
        # Load the Word template
        doc = Document(template_path)

        # Replace placeholders in all paragraphs
        for paragraph in doc.paragraphs:
            replace_placeholders_in_paragraph(paragraph, data)

        # Replace placeholders in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholders_in_paragraph(paragraph, data)

        # Save the filled document
        doc.save(output_path)
    except Exception as e:
        print(f"Error filling template: {e}")
        raise

def replace_placeholders_in_paragraph(paragraph, data):
    """
    Replace placeholders in a paragraph, accounting for multi-run issues.
    :param paragraph: Paragraph object.
    :param data: Dictionary containing placeholders and their values.
    """
    # Combine all runs into a single string
    full_text = ''.join(run.text for run in paragraph.runs)

    # Replace placeholders in the combined text
    for placeholder, value in data.items():
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, value)

    # Clear existing runs
    for run in paragraph.runs:
        run.text = ""

    # Write the updated text back into the paragraph
    if paragraph.runs:
        paragraph.runs[0].text = full_text

def generate_document(request_obj):
    # Path to the correct template based on document type
    templates = {
        "Barangay Clearance": os.path.join(settings.BASE_DIR, 'mambaling/templates/2024-barangay-clearance-final.docx'),
        "Certificate of Indigency": os.path.join(settings.BASE_DIR, 'mambaling/templates/barangay-indigency-written.docx'),
        "Certificate of Residency": os.path.join(settings.BASE_DIR, 'mambaling/templates/2024-CERT.-OF-RESIDENCY1.docx'),
    }
    template_path = templates.get(request_obj.document_type)

    # Load the template
    doc = Document(template_path)

    # Replace placeholders
    placeholders = {
        '{full_name}': request_obj.full_name,
        '{address}': request_obj.address,
        '{birthplace}': request_obj.birthplace,
        '{birthdate}': request_obj.birthdate.strftime('%B %d, %Y'),
        '{purpose}': request_obj.purpose or '',
        '{duration_of_residency}': request_obj.residency_duration or '',
    }
    for paragraph in doc.paragraphs:
        for key, value in placeholders.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    # Save the modified `.docx`
    output_docx_path = os.path.join(settings.MEDIA_ROOT, f"documents/{request_obj.id}_document.docx")
    doc.save(output_docx_path)

    # Convert to PDF (you can integrate this with a library like `pypandoc` or an external tool)
    output_pdf_path = os.path.join(settings.MEDIA_ROOT, f"documents/{request_obj.id}_document.pdf")
    return output_pdf_path
